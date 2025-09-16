import io
import os
import re
import zipfile
from typing import List, Tuple, Dict, Any, Optional

from fastapi import FastAPI, File, UploadFile, Header, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel

from docx import Document
import pandas as pd

# =========================
# Configuración básica
# =========================
MAX_DOCX = 10

app = FastAPI(title="Validador de Transcripciones DOCX vs Excel")

# CORS: tu dominio en producción
ALLOWED_ORIGINS = [
    "https://www.dipli.ai",
    "https://dipli.ai",
    "https://isagarcivill09.wixsite.com/turop",
    "https://isagarcivill09.wixsite.com/turop/tienda"
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Token opcional (recomendado en producción)
API_TOKEN: Optional[str] = os.getenv("API_TOKEN", None)

def check_token(x_token: Optional[str]):
    # Si no configuras API_TOKEN en el entorno, no se valida (modo dev)
    if API_TOKEN is None:
        return
    if x_token != API_TOKEN:
        raise HTTPException(status_code=401, detail="Unauthorized")

# =========================
# Reglas de validación
# =========================
ETIQUETAS_VALIDAS = ["ENTREVISTADOR", "ENTREVISTADORA", "ENTREVISTADO", "ENTREVISTADA"]
ETIQUETAS_INVALIDAS_PATRONES = [
    r"\bUsuario\b",
    r"\bX{3,}\b",
    r"\bXxxxxxx\b",
    r"\bEntrevistador[a-z]*\b",
    r"\bEntrevistado[a-z]*\b",
    r"\bentrevistador[a-z]*\b",
    r"\bentrevistado[a-z]*\b",
    r"\bModerador[a-z]*\b",
    r"\bmoderador[a-z]*\b",
    r"\bSpeaker[a-z]*\b",
    r"\bSPEAKER[a-z]*\b",
]

# =========================
# Utilidades
# =========================
def read_excel_questions(excel_bytes: bytes) -> List[str]:
    """Lee la columna C del .xlsx como matriz de referencia."""
    df = pd.read_excel(io.BytesIO(excel_bytes), usecols="C", engine="openpyxl").dropna()
    return df.iloc[:, 0].astype(str).tolist()

def doc_from_bytes(docx_bytes: bytes) -> Document:
    return Document(io.BytesIO(docx_bytes))

def validar_negrita_entrevistador(doc: Document) -> List[Tuple[int, str, str]]:
    errores = []
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        if texto.startswith("ENTREVISTADOR:") or texto.startswith("ENTREVISTADORA:"):
            is_bold = True
            for run in para.runs:
                if run.text.strip() and not run.bold:
                    is_bold = False
                    break
            if not is_bold:
                errores.append((
                    i + 1,
                    "Formato en negrita",
                    f"Texto de {texto.split(':')[0]} no está completamente en negrita."
                ))
    return errores

def validar_etiquetas_docx(doc: Document) -> List[Tuple[int, str, str]]:
    errores = []
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        match = re.match(r"^([A-Z]+):", texto)
        if match:
            etiqueta = match.group(1)
            if etiqueta not in ETIQUETAS_VALIDAS:
                errores.append((i + 1, "Etiqueta inválida", f"Etiqueta '{etiqueta}' no es válida."))
            elif etiqueta in ["ENTREVISTADOR", "ENTREVISTADORA"]:
                # Debe existir una run en negrita que contenga la etiqueta
                if not any(run.bold and etiqueta in run.text for run in para.runs):
                    errores.append((i + 1, "Encabezado sin negrita", f"Etiqueta '{etiqueta}' debería estar en negrita."))
    return errores

def detectar_etiquetas_invalidas(doc: Document) -> List[Tuple[int, str, str]]:
    errores = []
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        for patron in ETIQUETAS_INVALIDAS_PATRONES:
            if re.search(patron, texto, re.IGNORECASE):
                errores.append((i + 1, "Etiqueta inválida", f"Se encontró coincidencia con patrón: {patron}"))
                break
    return errores

def extraer_preguntas_docx(doc: Document) -> List[Tuple[int, str]]:
    """Extrae oraciones que terminan con '?' (incluye español)."""
    patron_pregunta = re.compile(r"[^?\n]+\?")
    preguntas = []
    for i, para in enumerate(doc.paragraphs):
        matches = patron_pregunta.findall(para.text)
        for pregunta in matches:
            pregunta_limpia = pregunta.strip()
            if pregunta_limpia:
                preguntas.append((i + 1, pregunta_limpia))
    return preguntas

def comparar_preguntas(preguntas_docx: List[Tuple[int, str]], preguntas_ref: List[str]) -> List[Tuple[int, str, str]]:
    errores = []
    for linea, pregunta in preguntas_docx:
        if pregunta not in preguntas_ref:
            errores.append((linea, "Pregunta no coincide", f"'{pregunta}' no está en la matriz de referencia"))
    return errores

def validar_fuente_tamano(doc: Document) -> List[Tuple[int, str, str]]:
    errores = []
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            fuente = run.font.name
            tam = run.font.size.pt if run.font.size else None  # .pt puede ser None
            if fuente and fuente.lower() != "arial":
                errores.append((i + 1, "Fuente incorrecta", f"'{fuente}' en vez de Arial"))
                break
            if tam and tam != 12:
                errores.append((i + 1, "Tamaño incorrecto", f"{tam}pt en vez de 12pt"))
                break
    return errores

def validar_intervencion_vacia(doc: Document) -> List[Tuple[int, str, str]]:
    errores = []
    for i, para in enumerate(doc.paragraphs):
        if re.match(r"^(ENTREVISTAD[OA]|ENTREVISTADOR[OA]):\s*$", para.text.strip()):
            errores.append((i + 1, "Intervención vacía tras etiqueta", "Debe ir junto al texto, sin salto de línea"))
    return errores

def analizar_un_docx(nombre: str, docx_bytes: bytes, preguntas_ref: List[str]) -> Dict[str, Any]:
    doc = doc_from_bytes(docx_bytes)

    errores: List[Tuple[int, str, str]] = []
    errores += comparar_preguntas(extraer_preguntas_docx(doc), preguntas_ref)
    errores += validar_etiquetas_docx(doc)
    errores += validar_negrita_entrevistador(doc)
    errores += validar_fuente_tamano(doc)
    errores += validar_intervencion_vacia(doc)
    errores += detectar_etiquetas_invalidas(doc)

    errores_list = [{"linea": l, "tipo_error": t, "descripcion": d} for (l, t, d) in errores]
    resumen: Dict[str, int] = {}
    for e in errores_list:
        resumen[e["tipo_error"]] = resumen.get(e["tipo_error"], 0) + 1

    return {"archivo": nombre, "resumen": resumen, "errores": errores_list}

# =========================
# Modelos de respuesta
# =========================
class AnalyzeResponse(BaseModel):
    ok: bool
    total_archivos: int
    resultados: List[Dict[str, Any]]

# =========================
# Endpoints
# =========================
@app.post("/analyze", response_model=AnalyzeResponse)
async def analyze_endpoint(
    documentos: List[UploadFile] = File(..., description="Máximo 10 archivos .docx"),
    preguntas_excel: UploadFile = File(..., description="Archivo .xlsx con matriz de preguntas (columna C)"),
    x_token: Optional[str] = Header(default=None)
):
    check_token(x_token)

    docx_files = [f for f in documentos if f.filename.lower().endswith(".docx")]
    if len(docx_files) == 0:
        return JSONResponse(status_code=400, content={"ok": False, "message": "Debes adjuntar al menos un .docx"})
    if len(docx_files) > MAX_DOCX:
        docx_files = docx_files[:MAX_DOCX]

    if not preguntas_excel.filename.lower().endswith(".xlsx"):
        return JSONResponse(status_code=400, content={"ok": False, "message": "Debes adjuntar un .xlsx de preguntas"})

    preguntas_ref = read_excel_questions(await preguntas_excel.read())

    resultados = []
    for f in docx_files:
        contenido = await f.read()
        resultados.append(analizar_un_docx(f.filename, contenido, preguntas_ref))

    return AnalyzeResponse(ok=True, total_archivos=len(resultados), resultados=resultados)

@app.post("/analyze-zip")
async def analyze_zip_endpoint(
    documentos: List[UploadFile] = File(..., description="Máximo 10 archivos .docx"),
    preguntas_excel: UploadFile = File(..., description="Archivo .xlsx con matriz de preguntas (columna C)"),
    x_token: Optional[str] = Header(default=None)
):
    check_token(x_token)

    docx_files = [f for f in documentos if f.filename.lower().endswith(".docx")]
    if len(docx_files) == 0:
        return JSONResponse(status_code=400, content={"ok": False, "message": "Debes adjuntar al menos un .docx"})
    if len(docx_files) > MAX_DOCX:
        docx_files = docx_files[:MAX_DOCX]

    if not preguntas_excel.filename.lower().endswith(".xlsx"):
        return JSONResponse(status_code=400, content={"ok": False, "message": "Debes adjuntar un .xlsx de preguntas"})

    preguntas_ref = read_excel_questions(await preguntas_excel.read())

    mem_zip = io.BytesIO()
    with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zipf:
        for f in docx_files:
            contenido = await f.read()
            resultado = analizar_un_docx(f.filename, contenido, preguntas_ref)

            base = os.path.splitext(os.path.basename(f.filename))[0]
            errores_txt = io.StringIO()
            resumen_txt = io.StringIO()

            # Detalle de errores
            if resultado["errores"]:
                for e in resultado["errores"]:
                    errores_txt.write(f"Línea {e['linea']}: {e['tipo_error']} - {e['descripcion']}\n")
            else:
                errores_txt.write("Sin errores detectados.\n")

            # Resumen por tipo
            if resultado["resumen"]:
                for tipo, count in resultado["resumen"].items():
                    resumen_txt.write(f"{tipo}: {count} ocurrencias\n")
            else:
                resumen_txt.write("Sin errores que resumir.\n")

            zipf.writestr(f"{base}_errores.txt", errores_txt.getvalue())
            zipf.writestr(f"{base}_resumen.txt", resumen_txt.getvalue())

    mem_zip.seek(0)
    headers = {"Content-Disposition": "attachment; filename=reportes_transcripciones.zip"}
    return StreamingResponse(mem_zip, media_type="application/zip", headers=headers)

@app.get("/health")
def health():
    return {"ok": True, "service": "validator", "version": "1.0.0"}
